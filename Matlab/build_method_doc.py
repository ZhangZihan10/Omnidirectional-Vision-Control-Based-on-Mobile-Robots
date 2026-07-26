from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT = "Method_Section_CN.docx"


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    for r in p.runs:
        r.font.name = "Microsoft YaHei"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r.font.size = Pt(9)


def set_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_code_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9.5)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.28)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.5)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.5)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        if level == 1:
            run.font.size = Pt(15)
            run.font.color.rgb = RGBColor(31, 78, 121)
        elif level == 2:
            run.font.size = Pt(12.5)
            run.font.color.rgb = RGBColor(46, 116, 181)
        else:
            run.font.size = Pt(11.5)
            run.font.color.rgb = RGBColor(31, 78, 121)
    return p


def build_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.9)
    sec.bottom_margin = Inches(0.9)
    sec.left_margin = Inches(0.9)
    sec.right_margin = Inches(0.9)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run("Method 部分草稿（中文）")
    r.bold = True
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(12)
    r = subtitle.add_run("基于 MATLAB-Unity 联合仿真的麦克纳姆轮小车全向视觉结构光 Mapping 与模糊自适应 PID 控制")
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor(90, 90, 90)

    add_heading(doc, "3 方法", 1)
    add_body(
        doc,
        "本文方法面向麦克纳姆轮移动机器人在虚拟环境中的实时感知、动态避障与运动控制问题。系统以 Unity 构建虚拟机器人和环境，"
        "以 MATLAB 完成全向视觉结构光测距、障碍物地图更新、路径规划和电机控制计算。整体方法由两个核心模块组成："
        "其一是基于全向相机与激光平面的实时 mapping 方法，用于将图像中的激光像素转换为环境中的二维障碍物坐标；"
        "其二是面向四个麦轮电机的前馈-模糊自适应 PID 速度控制方法，用于保证规划路径能够被稳定执行。"
    )

    add_heading(doc, "3.1 方法总体框架", 2)
    add_body(
        doc,
        "联合仿真闭环可以表示为“感知-建图-规划-控制-反馈”的连续过程。Unity 端提供全向相机图像和虚拟环境，"
        "MATLAB 端首先提取激光结构光区域，并依据全向相机标定模型和激光平面外参完成测距映射；随后将测距点投影到栅格地图，"
        "完成障碍物确认、地图衰减和剩余路径碰撞检测。当检测到当前路径与障碍物存在冲突时，系统触发局部 RRT* 重规划，"
        "并将原始路径转换为麦克纳姆轮小车可执行的八方向路径。最后，控制器根据当前路径点生成车体速度指令，"
        "通过麦轮逆运动学得到四轮目标速度，并由前馈-模糊自适应 PID 输出四轮 PWM 指令。"
    )
    add_code_paragraph(
        doc,
        "Unity 图像 → 激光提取 → 全向结构光 mapping → 栅格地图 → 局部重规划 → 车体速度规划 → 麦轮逆运动学 → 模糊 PID → 位置反馈"
    )

    add_heading(doc, "3.2 全向视觉与结构光 Mapping 模型", 2)
    add_body(
        doc,
        "本文采用全向相机与激光平面组成结构光测距系统。参考全向视觉结构光系统的建模思想，图像中的激光点可以看作相机视线"
        "与激光平面的交点。与传统针孔相机不同，全向或鱼眼相机需要通过标定模型将像素点反投影为空间射线。本文使用 OCamCalib"
        " 标定模型提供的 cam2world 函数完成该反投影。对于激光二值图中任意非零像素 m=[u,v]^T，有"
    )
    add_code_paragraph(doc, "M = cam2world(m, ocam_model),    M = [Mx, My, Mz]^T")
    add_body(
        doc,
        "其中 M 表示该像素对应的相机坐标系下空间射线方向。激光平面由其相对于相机的姿态和距离确定。程序中激光平面姿态由"
        " lasX、lasY 给出，激光平面到相机原点的距离由 las_dist 给出；相机自身安装姿态由 camX、camY、camZ 描述。"
        "设激光平面局部坐标中的一点为 P_L=[X,Y,1]^T，则其在相机坐标系中的表示可写为"
    )
    add_code_paragraph(doc, "P_C = H P_L,    H = R_C [ r1  r2  t ],    t = [0, 0, las_dist]^T")
    add_body(
        doc,
        "其中 R_C=R(camX,camY,camZ) 为相机安装旋转矩阵，[r1,r2] 由激光平面旋转 R(-lasX,-lasY,0) 的前两列构成。"
        "由于图像射线与激光平面交于同一个空间点，因此存在比例因子 λ，使得"
    )
    add_code_paragraph(doc, "λ M = H [X, Y, 1]^T")
    add_body(
        doc,
        "对比例因子 λ 进行消元后，可以得到关于 X 和 Y 的两个线性方程。程序中的 mapping 函数将其写为"
    )
    add_code_paragraph(
        doc,
        "a1 X + b1 Y + c1 = 0\n"
        "a2 X + b2 Y + c2 = 0\n"
        "Y = (a2 c1 - a1 c2) / (a1 b2 - a2 b1),    X = (-c1 - b1 Y) / a1"
    )
    add_body(
        doc,
        "由此，每个激光像素都可以被转换为激光平面上的二维交点。随后，系统根据小车或视觉系统的当前姿态对该点进行旋转和平移，"
        "得到全局坐标系中的测距点："
    )
    add_code_paragraph(doc, "P_W = R_z(-CVsyst_rot) [X, Y, 1]^T + T_CV")
    add_body(
        doc,
        "在具体实现中，由于坐标轴定义与 MATLAB 图像坐标方向存在差异，程序对 X、Y 分量进行了符号和轴向变换。"
        "最终输出的点集被用作环境障碍物点云，并进一步参与栅格地图更新。"
    )

    add_heading(doc, "3.3 激光提取与实时障碍物地图", 2)
    add_body(
        doc,
        "在虚拟环境中，激光提取通过颜色阈值获得二值激光区域；在真实相机测试程序中，则采用 HSV 红色范围与 RGB 红色通道优势"
        "相结合的方式抑制背景噪声。提取得到的激光区域不再必须骨架化，而是保留经过连通域过滤后的完整激光区域，"
        "从而在激光条纹较弱或存在断裂时保留更多有效测距点。"
    )
    add_body(
        doc,
        "映射后的点云会被投影到二维栅格地图中。每个栅格保存命中次数和最近观测帧号，只有命中次数超过阈值的栅格才被认定为"
        "确认障碍物。该策略可以降低单帧噪声对路径规划的影响。为了适应动态环境，系统还引入地图衰减机制：若某个栅格在一段时间内"
        "没有被再次观测，其可信度逐渐降低，最终从确认障碍物集合中移除。"
    )
    add_body(
        doc,
        "路径监测只检查小车当前位置之后的剩余路径，而不检查已经走过的历史路径。这样可以避免车后方障碍物引发不必要的重规划。"
        "当剩余路径线段与确认障碍物距离小于安全阈值时，系统触发局部重规划，并重新生成麦轮可执行路径。"
    )

    add_heading(doc, "3.4 麦克纳姆轮速度生成与逆运动学", 2)
    add_body(
        doc,
        "路径规划输出一系列离散路径点。控制器首先计算当前位置到当前目标路径点的误差向量，并根据距离设计目标车体速度。"
        "当小车距离目标路径点较远时采用较高速度，当接近目标路径点时逐渐减速，以降低越点和转向冲击。为了避免低帧率条件下速度指令突变，"
        "系统对车体速度变化率设置上限。"
    )
    add_code_paragraph(
        doc,
        "v_ref = direction · v_target\n"
        "||v_ref(k) - v_ref(k-1)|| <= min(MAX_CHASSIS_ACCEL · dt, MAX_CHASSIS_REF_STEP)"
    )
    add_body(
        doc,
        "得到车体速度 vx、vy 和角速度 wz 后，通过麦克纳姆轮逆运动学计算四个轮子的目标线速度。程序采用的轮速分配关系为"
    )
    add_code_paragraph(
        doc,
        "v0 = vy - vx + kω wz\n"
        "v1 = vy + vx - kω wz\n"
        "v2 = vy - vx - kω wz\n"
        "v3 = vy + vx + kω wz"
    )
    add_body(
        doc,
        "其中 v0 至 v3 分别表示四个麦轮的目标线速度，kω 为旋转速度折算系数。当前仿真主要验证平移运动控制，因此通常令 wz=0。"
    )

    add_heading(doc, "3.5 前馈-模糊自适应 PID 电机控制", 2)
    add_body(
        doc,
        "四个电机分别采用独立的速度闭环控制器。为了提高响应速度，系统首先根据目标轮速计算前馈 PWM，作为主要驱动力："
    )
    add_code_paragraph(doc, "PWM_ff = v_ref / MOTOR_GAIN")
    add_body(
        doc,
        "随后，模糊自适应 PID 根据目标轮速与实际轮速之间的误差计算反馈修正量。设第 i 个电机的目标轮速为 r，实际轮速为 y，"
        "则误差和误差变化率为"
    )
    add_code_paragraph(doc, "e = r - y,    ec = (e(k) - e(k-1)) / dt")
    add_body(
        doc,
        "模糊自适应部分并不采用固定 PID 参数，而是根据误差幅值、误差变化幅值以及误差是否继续增大来动态调整 Kp、Ki、Kd。"
        "在代码中，误差归一化变量为"
    )
    add_code_paragraph(
        doc,
        "e_level = min(|e| / E_MAX, 1)\n"
        "ec_level = min(|ec| / EC_MAX, 1)\n"
        "g = 0.5 · (1 + tanh(e · ec / 80000))"
    )
    add_body(
        doc,
        "其中 g 用于描述误差是否有继续扩大的趋势。若 e 与 ec 同号，则误差正在变大，控制器需要更强的比例和阻尼作用。"
        "PID 参数调度规则为"
    )
    add_code_paragraph(
        doc,
        "Kp = 0.022 + 0.026 e_level + 0.008 g\n"
        "Ki = 0.0010 + 0.0060 (1 - e_level)^2\n"
        "Kd = 0.0020 + 0.0050 ec_level + 0.0025 (1 - e_level)"
    )
    add_body(
        doc,
        "上述规则体现了模糊 PID 的基本思想：当误差较大时增大 Kp，提高响应速度；当误差较小时增大 Ki，以消除稳态误差；"
        "当误差变化较快时增大 Kd，以抑制超调和振荡。最终的 PID 修正量为"
    )
    add_code_paragraph(doc, "PWM_PID = Kp e + Ki ∫e dt + Kd d(e)/dt")
    add_body(
        doc,
        "为了避免目标轮速突变导致微分冲击，程序采用对实际轮速求导的方式构造微分项，即 derivative on measurement。"
        "同时使用一阶滤波平滑微分信号。最终电机控制指令为"
    )
    add_code_paragraph(doc, "PWM_cmd = sat(PWM_ff + PWM_PID, -PWM_max, PWM_max)")
    add_body(
        doc,
        "其中 sat 表示限幅函数。为提高低帧率仿真下的稳定性，系统进一步限制单帧 PWM 变化量，并根据实际控制周期自动降低 PID 修正上限。"
        "此外，当轮速指令发生明显方向反转时，系统清除上一方向累积的积分项，防止旧积分与新方向控制相互抵消。"
    )

    add_heading(doc, "3.6 控制器稳定性保护机制", 2)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    set_cell_text(hdr[0], "机制", True)
    set_cell_text(hdr[1], "作用", True)
    for c in hdr:
        set_shading(c, "E8EEF5")
    rows = [
        ("前馈控制", "根据目标轮速直接给出基础 PWM，减小 PID 需要承担的误差修正量。"),
        ("积分分离", "误差较大时不继续累积积分，避免大误差阶段产生积分饱和。"),
        ("条件积分", "当输出已经饱和且积分会进一步推动饱和时，暂停积分更新。"),
        ("方向反转积分清零", "轮速目标正负方向发生明显切换时，清除旧方向积分。"),
        ("微分滤波", "降低实际轮速测量变化造成的高频抖动。"),
        ("真实时间步长", "用实际循环时间 dt 进行 PID、电机模型和位置积分，避免帧率变化引起控制尺度错误。"),
        ("PWM 变化率限制", "限制单帧 PWM 突变，降低转向和动态重规划时的速度冲击。"),
    ]
    for left, right in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], left)
        set_cell_text(cells[1], right)

    add_heading(doc, "3.7 方法小结", 2)
    add_body(
        doc,
        "综上，本文方法将全向视觉结构光测距、动态地图更新、路径重规划和四轮电机控制统一到同一个 MATLAB-Unity 闭环仿真框架中。"
        "全向视觉结构光模块提供环境点云和障碍物位置，路径规划模块生成满足麦克纳姆轮运动特性的可执行路径，"
        "前馈-模糊自适应 PID 控制器则保证小车能够在低帧率和动态环境条件下稳定跟踪目标轮速。该方法为后续真实小车平台移植提供了"
        "完整的感知、规划和控制算法基础。"
    )

    add_heading(doc, "参考文献占位", 1)
    add_body(
        doc,
        "[1] I. Y. Kholodilin, Y. Li, and Q. Wang, “Omnidirectional Vision System with Laser Illumination in a Flexible Configuration and its Calibration by One Single Snapshot,” IEEE Transactions on Instrumentation and Measurement, 2020."
    )
    add_body(
        doc,
        "[2] 本文 MATLAB 程序：ARealTimeTest8_3.m, mapping.m, fuzzyAdaptivePIDCorrection3.m, initFuzzyPIDState3.m。"
    )

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
